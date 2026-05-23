from __future__ import annotations

import csv
import json
import subprocess
import ast
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
BASE_REPORT = ROOT / "report_lab2_v16_defense.docx"
OUTPUT_REPORT = ROOT / "report_lab2_v16_final.docx"
FORMULA_DIR = ROOT / "formula_images" / "lab2_final"

FONT = "Times New Roman"
MONO = "Consolas"
BODY_SIZE = Pt(14)
TABLE_SIZE = Pt(11)
SMALL_TABLE_SIZE = Pt(10.5)

plt.rcParams.update({"mathtext.fontset": "stix", "font.family": "STIXGeneral"})


def set_run_font(run, *, bold: bool = False, italic: bool = False, size=BODY_SIZE, mono: bool = False) -> None:
    name = MONO if mono else FONT
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.different_first_page_header_footer = True

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = BODY_SIZE
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run)


def add_p(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    align=None,
    indent: bool = True,
    size=BODY_SIZE,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size=size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    for run in paragraph.runs:
        set_run_font(run, bold=True, size=BODY_SIZE)


def add_dash(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(-0.6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run("- " + text)
    set_run_font(run)


def add_code_block(doc: Document, code: str, *, size=Pt(7.3)) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shading)
    lines = code.splitlines()
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, mono=True, size=size)
        if idx < len(lines) - 1:
            run.add_break()


def extract_function_source(file_path: Path, func_name: str) -> str:
    source = file_path.read_text(encoding="utf-8")
    tree = ast.parse(source)
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == func_name:
            return "\n".join(source.splitlines()[node.lineno - 1 : node.end_lineno])
    raise ValueError(f"{func_name} not found in {file_path}")


def page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_title_page(doc: Document) -> None:
    for line in [
        "Министерство образования и науки Российской Федерации",
        "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        "«Ижевский государственный технический университет имени М. Т. Калашникова»",
    ]:
        add_p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(3):
        add_p(doc, "", indent=False)

    add_p(doc, "Лабораторная работа N 2", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
    add_p(doc, "Решение задач линейного программирования", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
    add_p(doc, "По дисциплине «Методы оптимизации»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_p(doc, "Вариант 16", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(2):
        add_p(doc, "", indent=False)

    add_p(doc, "Выполнил: студент гр. М25-787-1 Р. В. Скороходов", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    add_p(doc, "Принял: доктор физико-математических наук, профессор В. А. Тененев", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)

    for _ in range(1):
        add_p(doc, "", indent=False)
    add_p(doc, "Ижевск 2026", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    page_break(doc)


def add_toc(doc: Document) -> None:
    add_p(doc, "Содержание", bold=True, indent=False)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Содержание будет обновлено автоматически."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    page_break(doc)


def render_formula(name: str, latex: str, *, width: float = 6.0, height: float = 0.75, fontsize: int = 16) -> Path:
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    path = FORMULA_DIR / f"{name}.png"
    fig = plt.figure(figsize=(width, height), dpi=300)
    fig.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=fontsize)
    fig.savefig(path, transparent=True, bbox_inches="tight", pad_inches=0.03)
    plt.close(fig)
    return path


def add_formula(doc: Document, latex: str, number: str, *, name: str, max_width_cm: float = 12.0, height: float = 0.75) -> None:
    path = render_formula(name, latex, width=max_width_cm / 2.2, height=height)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t")
    run = paragraph.add_run()
    with Image.open(path) as image:
        dpi = image.info.get("dpi", (300, 300))[0] or 300
        native_width_cm = image.size[0] / dpi * 2.54
    if native_width_cm > max_width_cm:
        run.add_picture(str(path), width=Cm(max_width_cm))
    else:
        run.add_picture(str(path))
    paragraph.add_run("\t")
    number_run = paragraph.add_run(f"({number})")
    set_run_font(number_run)


def set_cell_text(cell, text: str, *, bold: bool = False, size=TABLE_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = align
    run = paragraph.add_run(str(text))
    set_run_font(run, bold=bold, size=size)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], *, size=TABLE_SIZE) -> None:
    add_p(doc, caption, bold=True, indent=False)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=size)
    add_p(doc, "", indent=False)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(13.2))
    add_p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as file:
        return list(csv.DictReader(file))


def load_data() -> dict:
    return json.loads((ROOT / "tables" / "summary.json").read_text(encoding="utf-8"))


def phase_rows(task: str, phase: str) -> list[list[str]]:
    rows = read_csv(ROOT / "tables" / f"simplex_{task}_steps.csv")
    return [
        [
            row["iteration"],
            row["basis"],
            row["rhs"],
            row["objective"],
            row["entering"] or "-",
            row["leaving"] or "-",
        ]
        for row in rows
        if row["phase"] == phase
    ]


def tableau_rows(task: str) -> tuple[list[str], list[list[str]]]:
    path = ROOT / "tables" / f"simplex_{task}_final_tableau.csv"
    rows = read_csv(path)
    headers = list(rows[0].keys())
    return headers, [[row[header] for header in headers] for row in rows]


def validate_results(summary: dict) -> None:
    assert summary["objective"] == "Z = -3*x1 + 6*x2"
    assert summary["min"]["point_fraction"] == ["12/7", "16/7"]
    assert summary["min"]["value_fraction"] == "60/7"
    assert summary["max"]["value_fraction"] == "12"
    vertices = {(item["x1_fraction"], item["x2_fraction"]) for item in summary["vertices"]}
    assert vertices == {("4/3", "8/3"), ("12/7", "16/7"), ("2", "3")}


def collect_doc_text(doc: Document) -> str:
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def update_fields_and_get_pages(path: Path) -> int:
    ps_path = str(path).replace("'", "''")
    command = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{ps_path}')
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
$doc.Fields.Update() | Out-Null
$pages = $doc.ComputeStatistics(2)
$doc.Save()
$doc.Close($false)
$word.Quit()
Write-Output $pages
"""
    result = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", command],
        check=True,
        capture_output=True,
        text=True,
    )
    return int(result.stdout.strip().splitlines()[-1])


def build_report() -> None:
    if not BASE_REPORT.exists():
        raise FileNotFoundError(BASE_REPORT)
    summary = load_data()
    validate_results(summary)

    doc = Document()
    setup_document(doc)
    add_title_page(doc)
    add_toc(doc)

    add_heading(doc, "1 Обозначения и ключевые понятия")
    add_table(
        doc,
        "Таблица 1.1. Обозначения и ключевые понятия",
        ["Обозначение", "Что означает"],
        [
            ["x = (x1, x2)", "Точка на плоскости"],
            ["Z", "Целевая функция"],
            ["Zmin", "Минимальное значение целевой функции"],
            ["Zmax", "Максимальное значение целевой функции"],
            ["x*", "Оптимальная точка или множество оптимальных точек"],
            ["s", "Добавочная переменная для ограничения типа <="],
            ["e", "Избыточная переменная для ограничения типа >="],
            ["a", "Искусственная переменная для построения начального базиса"],
            ["базис", "Набор переменных, через которые выражается текущее решение"],
            ["ведущий столбец", "Столбец переменной, которая входит в базис"],
            ["ведущая строка", "Строка переменной, которая выходит из базиса"],
            ["оценочная строка", "Строка, по которой проверяется возможность улучшения Z"],
        ],
        size=Pt(11),
    )

    add_heading(doc, "1.1 Зачем нужна эта работа", level=2)
    add_p(
        doc,
        "Линейное программирование используется, когда нужно выбрать лучший вариант при ограниченных ресурсах. "
        "Например, можно максимизировать прибыль, минимизировать затраты, распределить сырье, время или транспорт. "
        "В этой лабораторной рассматривается задача с двумя переменными, поэтому ее можно решить графически. "
        "Это удобно: на рисунке видно допустимую область и точки, где функция принимает минимум и максимум. "
        "Затем тот же результат проверяется симплекс-методом, который работает уже не по рисунку, а через таблицы.",
    )

    add_heading(doc, "2 Цель работы")
    add_p(
        doc,
        "Цель работы - решить задачу линейного программирования графическим методом и методом симплекс-таблиц, "
        "найти минимум и максимум целевой функции и проверить совпадение результатов двумя способами.",
    )

    add_heading(doc, "3 Постановка задачи")
    add_p(
        doc,
        "Дана целевая функция Z = -3*x1 + 6*x2 и система линейных ограничений. Требуется найти минимальное "
        "и максимальное значения Z при выполнении всех ограничений. Переменные x1 и x2 неотрицательны. "
        "Задача относится к линейному программированию, потому что и целевая функция, и ограничения являются линейными.",
    )
    add_p(doc, "Формула целевой функции:")
    add_formula(doc, r"Z=-3x_1+6x_2", "3.1", name="objective", max_width_cm=7.5)
    add_p(doc, "Ограничения задачи:")
    add_formula(
        doc,
        r"\left\{\substack{5x_1-2x_2\leq4\\x_1-2x_2\geq-4\\x_1+x_2\geq4\\x_1\geq0\\x_2\geq0}\right.",
        "3.2",
        name="constraints",
        max_width_cm=8.2,
        height=1.25,
    )
    add_p(doc, "Требуется найти:")
    add_formula(doc, r"Z_{\min}=\min Z,\quad Z_{\max}=\max Z", "3.3", name="minmax", max_width_cm=9.2)
    add_table(
        doc,
        "Таблица 3.1. Ограничения задачи",
        ["N", "Ограничение", "Граничная прямая", "Полуплоскость"],
        [
            ["1", "5*x1 - 2*x2 <= 4", "5*x1 - 2*x2 = 4", "область с проверочной точкой (0; 0)"],
            ["2", "x1 - 2*x2 >= -4", "x1 - 2*x2 = -4", "область с проверочной точкой (0; 0)"],
            ["3", "x1 + x2 >= 4", "x1 + x2 = 4", "выше прямой"],
            ["4", "x1 >= 0", "x1 = 0", "правая полуплоскость"],
            ["5", "x2 >= 0", "x2 = 0", "верхняя полуплоскость"],
        ],
        size=Pt(10.5),
    )

    add_heading(doc, "4 Графический метод")
    add_p(
        doc,
        "Графический метод состоит из четырех шагов. Каждое ограничение заменяется граничной прямой. "
        "По знаку неравенства выбирается нужная полуплоскость. Пересечение всех полуплоскостей образует "
        "допустимую область. Затем значение Z проверяется в вершинах допустимой области.",
    )
    add_p(
        doc,
        "Для линейной функции минимум и максимум на многоугольнике достигаются в вершинах или на целом ребре. "
        "Поэтому достаточно найти вершины допустимой области и подставить их координаты в Z.",
    )
    add_table(
        doc,
        "Таблица 4.1. Значения целевой функции в вершинах допустимой области",
        ["Точка", "x1", "x2", "Z = -3*x1 + 6*x2", "Активные ограничения"],
        [
            ["A1", "4/3", "8/3", "12", "x1 - 2*x2 = -4; x1 + x2 = 4"],
            ["A2", "12/7", "16/7", "60/7", "5*x1 - 2*x2 = 4; x1 + x2 = 4"],
            ["A3", "2", "3", "12", "5*x1 - 2*x2 = 4; x1 - 2*x2 = -4"],
        ],
        size=Pt(11),
    )
    add_p(
        doc,
        "Минимальное значение получается в точке A2, потому что 60/7 < 12. Максимальное значение равно 12. "
        "Оно достигается сразу в двух соседних вершинах A1 и A3. Так как целевая функция линейная, все точки "
        "отрезка A1A3 также дают Z = 12. Поэтому максимум не единственный.",
    )
    add_figure(
        doc,
        ROOT / "figures" / "feasible_region.png",
        "Рисунок 4.1. Допустимая область задачи линейного программирования. Голубой многоугольник показывает "
        "все допустимые решения. Точка A2 дает минимум Z = 60/7, а оранжевое ребро A1A3 дает максимум Z = 12.",
    )
    add_p(doc, "Что нужно запомнить:", bold=True)
    add_dash(doc, "допустимая область получается как пересечение полуплоскостей;")
    add_dash(doc, "для линейной функции экстремум достигается в вершине или на ребре;")
    add_dash(doc, "если две соседние вершины дают одинаковое максимальное значение, максимум достигается на всем отрезке между ними.")

    add_heading(doc, "5 Симплекс-метод")
    add_p(
        doc,
        "Симплекс-таблица показывает текущее базисное решение. Столбец свободных членов содержит значения базисных "
        "переменных. Оценочная строка показывает, можно ли улучшить целевую функцию. Если для максимизации в "
        "оценочной строке есть положительный коэффициент, решение еще можно улучшить. Ведущий столбец показывает "
        "переменную, которая входит в базис. Ведущая строка показывает переменную, которая выходит из базиса.",
    )

    add_heading(doc, "5.1 Приведение задачи к каноническому виду", level=2)
    add_formula(
        doc,
        r"\left\{\substack{5x_1-2x_2+s_1=4\\-x_1+2x_2+s_2=4\\x_1+x_2-e_3+a_3=4}\right.",
        "5.1",
        name="canonical",
        max_width_cm=9.2,
        height=1.05,
    )
    add_p(doc, "Первое ограничение имеет тип <=, поэтому добавляется переменная s1.")
    add_p(doc, "Второе ограничение сначала умножается на -1, чтобы правая часть стала положительной, затем добавляется s2.")
    add_p(doc, "Третье ограничение имеет тип >=, поэтому из него вычитается e3 и добавляется искусственная переменная a3.")

    add_heading(doc, "5.2 Фаза I. Поиск начального допустимого базиса", level=2)
    add_formula(doc, r"F=-a_3\to\max", "5.2", name="phase1", max_width_cm=6.5)
    add_table(
        doc,
        "Таблица 5.1. Фаза I симплекс-метода",
        ["k", "Базис", "Свободные члены", "Целевая строка", "Входит", "Выходит"],
        phase_rows("max", "phase I"),
        size=Pt(11),
    )
    add_p(
        doc,
        "В фазе I искусственная переменная должна уйти из базиса. Если оптимальное значение вспомогательной функции "
        "равно нулю, допустимое решение найдено. В данной задаче после двух переходов получен базис x1, s2, x2, "
        "поэтому можно переходить к исходной целевой функции.",
    )

    add_heading(doc, "5.3 Фаза II. Поиск максимума", level=2)
    add_p(doc, "Для максимизации используется исходная функция:")
    add_formula(doc, r"Z=-3x_1+6x_2\to\max", "5.3", name="max_z", max_width_cm=8.5)
    add_p(
        doc,
        "После фазы I текущая допустимая вершина имеет значение Z = 60/7. Симплекс-метод проверяет, можно ли "
        "увеличить значение Z. После одного перехода получается точка (2; 3) и значение Z = 12.",
    )
    add_table(
        doc,
        "Таблица 5.2. Фаза II для максимума",
        ["k", "Базис", "Свободные члены", "Целевая строка", "Входит", "Выходит"],
        phase_rows("max", "phase II"),
        size=Pt(11),
    )
    headers, rows = tableau_rows("max")
    add_table(doc, "Таблица 5.3. Итоговая симплекс-таблица для максимума", headers, rows, size=Pt(11))
    add_p(
        doc,
        "Нулевая оценка у небазисной переменной означает альтернативный оптимум. Поэтому симплекс-метод дает одну "
        "вершину оптимального ребра, а графический метод показывает все множество решений: отрезок от (4/3; 8/3) до (2; 3).",
    )

    add_heading(doc, "5.4 Фаза II. Поиск минимума", level=2)
    add_p(
        doc,
        "Симплекс-алгоритм в работе реализован для максимизации. Поэтому минимум исходной функции находится через "
        "максимум противоположной функции:",
    )
    add_formula(doc, r"W=-Z=3x_1-6x_2\to\max", "5.4", name="min_w", max_width_cm=9.0)
    add_p(doc, "Если максимум W найден в некоторой точке, то в этой же точке исходная функция Z принимает минимум.")
    add_table(
        doc,
        "Таблица 5.4. Фаза II для минимума через максимизацию -Z",
        ["k", "Базис", "Свободные члены", "Целевая строка", "Входит", "Выходит"],
        phase_rows("min", "phase II"),
        size=Pt(11),
    )
    headers, rows = tableau_rows("min")
    add_table(doc, "Таблица 5.5. Итоговая симплекс-таблица для минимума", headers, rows, size=Pt(11))
    add_p(
        doc,
        "Положительных оценок для функции W = -Z нет уже на первой строке фазы II. Поэтому текущая вершина "
        "(12/7; 16/7) является оптимальной для задачи минимизации исходной функции. Значение исходной функции "
        "равно Zmin = 60/7.",
    )
    add_p(doc, "Что нужно запомнить:", bold=True)
    add_dash(doc, "симплекс-метод переходит от одного базисного решения к другому;")
    add_dash(doc, "фаза I нужна для поиска начального допустимого базиса;")
    add_dash(doc, "фаза II оптимизирует исходную целевую функцию;")
    add_dash(doc, "нулевая оценка у небазисной переменной может означать альтернативный оптимум.")

    add_heading(doc, "6 Контрольная проверка")
    add_p(
        doc,
        "Контрольная проверка используется только для подтверждения результата. Она не заменяет графический метод "
        "и симплекс-таблицы, потому что по заданию требуется показать именно эти два способа решения.",
    )
    add_table(
        doc,
        "Таблица 6.1. Сравнение результатов",
        ["Способ", "Минимум", "Где достигается минимум", "Максимум", "Где достигается максимум"],
        [
            ["Графический метод", "60/7", "(12/7; 16/7)", "12", "отрезок A1A3"],
            ["Симплекс-метод", "60/7", "(12/7; 16/7)", "12", "обнаружена вершина и признак альтернативного оптимума"],
            ["Контрольная проверка Python", "60/7", "(12/7; 16/7)", "12", "отрезок A1A3"],
        ],
        size=Pt(10.5),
    )

    add_heading(doc, "7 Сравнение графического и симплекс-метода")
    add_p(
        doc,
        "Графический метод удобен для задачи с двумя переменными, потому что допустимую область можно нарисовать "
        "на плоскости. Его главное преимущество - наглядность. На графике сразу видно, что максимум достигается на целом ребре.",
    )
    add_p(
        doc,
        "Симплекс-метод менее нагляден, но более универсален. Он работает через таблицы и может применяться к задачам "
        "с большим числом переменных, где график построить невозможно. В данной задаче оба метода дали одинаковые "
        "значения Zmin и Zmax, значит решение выполнено корректно.",
    )

    add_heading(doc, "8 Выводы")
    add_p(
        doc,
        "В работе задача линейного программирования решена графическим методом и методом симплекс-таблиц. "
        "Графический метод дал вершины допустимой области A1 = (4/3; 8/3), A2 = (12/7; 16/7), A3 = (2; 3). "
        "Минимум равен Zmin = 60/7 и достигается в точке A2. Максимум равен Zmax = 12 и достигается на отрезке "
        "от (4/3; 8/3) до (2; 3). Симплекс-метод подтвердил эти значения и выявил признак альтернативного оптимума "
        "для задачи максимизации.",
    )

    add_heading(doc, "9 Список использованных источников")
    add_p(doc, "Лабораторная работа N 2. Решение задач линейного программирования: методические указания.", indent=False)
    add_p(doc, "Репозиторий с исходным кодом и ноутбуками: https://github.com/megusto0/mp-2.", indent=False)

    page_break(doc)
    add_heading(doc, "Приложение А. Репозиторий и ноутбуки")
    add_p(
        doc,
        "Все вычисления, построение графика, поиск вершин допустимой области и формирование таблиц выполнены в Python. "
        "Репозиторий содержит ноутбуки, которые можно открыть в Google Colab и выполнить сверху вниз.",
    )
    add_p(doc, "Репозиторий: https://github.com/megusto0/mp-2", indent=False)
    add_p(doc, "Структура репозитория:", indent=False)
    for item in [
        "notebooks/01_graphical_method.ipynb - постановка задачи и графический метод;",
        "notebooks/02_simplex_max.ipynb - симплекс-метод для максимума;",
        "notebooks/03_simplex_min.ipynb - симплекс-метод для минимума;",
        "notebooks/04_summary.ipynb - сравнение результатов;",
        "src/lp2/problem.py - коэффициенты задачи;",
        "src/lp2/graphical.py - поиск вершин допустимой области;",
        "src/lp2/simplex.py - симплекс-таблицы;",
        "src/lp2/plotting.py - построение графиков;",
        "figures/ - сохраненные графики;",
        "tables/ - сохраненные таблицы.",
    ]:
        add_dash(doc, item)

    add_heading(doc, "А.1 Листинги основных функций", level=2)
    add_p(
        doc,
        "Ниже приведены оформленные листинги функций, которые отвечают за поиск вершин допустимой области, "
        "построение графика и выполнение симплекс-метода. В основной текст код не вынесен, чтобы не перегружать решение.",
    )
    for file_name, func_name in [
        ("src/lp2/graphical.py", "enumerate_vertices"),
        ("src/lp2/graphical.py", "solve_graphically"),
        ("src/lp2/simplex.py", "standardize"),
        ("src/lp2/simplex.py", "_run_simplex"),
        ("src/lp2/simplex.py", "solve_with_simplex"),
        ("src/lp2/plotting.py", "plot_feasible_region"),
    ]:
        add_p(doc, f"Функция {func_name} из файла {file_name}", bold=True, indent=False)
        add_code_block(doc, extract_function_source(ROOT / file_name, func_name))

    page_break(doc)
    add_heading(doc, "Приложение Б. Вопросы и ответы для защиты")
    qa = [
        ("Вопрос 1. Что такое допустимая область?", "Это множество всех точек (x1, x2), которые удовлетворяют всем ограничениям задачи."),
        ("Вопрос 2. Почему экстремум ищется в вершинах?", "Потому что целевая функция линейная. На выпуклом многоугольнике линейная функция достигает минимума или максимума в вершине или на целом ребре."),
        ("Вопрос 3. Почему максимум получился не в одной точке?", "В двух соседних вершинах A1 и A3 значение Z одинаковое и равно 12. Значит, на всем отрезке между этими вершинами значение Z тоже равно 12."),
        ("Вопрос 4. Зачем нужна искусственная переменная a3?", "Она нужна, чтобы построить начальный базис для ограничения типа >=. После фазы I искусственная переменная должна уйти из базиса."),
        ("Вопрос 5. Что делает фаза I симплекс-метода?", "Фаза I ищет начальное допустимое базисное решение. Если вспомогательная функция стала равна нулю, допустимый базис найден."),
        ("Вопрос 6. Что делает фаза II?", "Фаза II уже оптимизирует исходную целевую функцию Z или противоположную функцию -Z, если ищется минимум."),
        ("Вопрос 7. Почему для минимума используется -Z?", "Потому что используемый симплекс-алгоритм настроен на максимизацию. Минимизация Z равносильна максимизации -Z."),
        ("Вопрос 8. Что значит нулевая оценка у небазисной переменной?", "Это признак альтернативного оптимума. Можно перейти в соседнюю вершину, не изменив значение целевой функции."),
        ("Вопрос 9. Почему графический и симплекс-метод должны совпасть?", "Они решают одну и ту же задачу. Графический метод делает это через рисунок, а симплекс-метод через таблицы."),
        ("Вопрос 10. Какой окончательный ответ?", "Zmin = 60/7 в точке (12/7; 16/7). Zmax = 12 на отрезке от (4/3; 8/3) до (2; 3)."),
    ]
    for question, answer in qa:
        add_p(doc, question, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
        add_p(doc, answer)

    doc_text = collect_doc_text(doc)
    assert chr(0x2014) not in doc_text and chr(0x2013) not in doc_text
    assert "очевидно" not in doc_text.lower()
    assert "тривиально" not in doc_text.lower()
    assert "легко видеть" not in doc_text.lower()

    doc.save(OUTPUT_REPORT)
    pages = update_fields_and_get_pages(OUTPUT_REPORT)
    final_doc = Document(OUTPUT_REPORT)
    final_text = collect_doc_text(final_doc)
    assert chr(0x2014) not in final_text and chr(0x2013) not in final_text
    assert pages <= 16, f"Report is too long: {pages} pages"
    print(f"Saved {OUTPUT_REPORT}")
    print(f"Pages: {pages}")


if __name__ == "__main__":
    build_report()
