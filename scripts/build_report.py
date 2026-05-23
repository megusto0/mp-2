from __future__ import annotations

import sys
import subprocess
from fractions import Fraction
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from src.lp2.graphical import format_fraction, solve_graphically
from src.lp2.problem import VARIANT_16
from src.lp2.simplex import solve_with_simplex

REPORT_PATH = ROOT / "report_lab2_v16_defense.docx"
FORMULA_DIR = ROOT / "formula_images"
FONT = "Times New Roman"
TEXT_SIZE = Pt(14)


def as_fraction(value: float) -> str:
    fraction = Fraction(float(value)).limit_denominator(1000)
    if abs(float(fraction) - float(value)) <= 1e-8:
        if fraction.denominator == 1:
            return str(fraction.numerator)
        return f"{fraction.numerator}/{fraction.denominator}"
    return f"{float(value):.4f}"


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = TEXT_SIZE

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = TEXT_SIZE

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run)


def set_run_font(run, *, bold: bool = False, italic: bool = False, size=TEXT_SIZE) -> None:
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)


def page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, align=None, indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, bold=True)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, df: pd.DataFrame, title: str, *, size: int = 9) -> None:
    add_paragraph(doc, title, bold=True, indent=False)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, column in enumerate(df.columns):
        set_cell_text(table.rows[0].cells[idx], str(column), bold=True, size=size)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(df.columns):
            set_cell_text(cells[idx], str(row[column]), size=size)
    doc.add_paragraph()


def render_math(lines: list[str], path: Path, *, brace: bool = False, width: float = 5.2) -> None:
    FORMULA_DIR.mkdir(exist_ok=True)
    height = max(0.45 * len(lines), 0.65)
    fig = plt.figure(figsize=(width, height), dpi=240)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    y_positions = list(reversed([(idx + 0.5) / len(lines) for idx in range(len(lines))]))
    x_text = 0.15 if brace else 0.02
    if brace:
        ax.text(0.03, 0.5, "{", fontsize=24 + 12 * len(lines), va="center", ha="left", family="DejaVu Serif")
    for y, line in zip(y_positions, lines):
        ax.text(x_text, y, f"${line}$", fontsize=16, va="center", ha="left", color="black")
    fig.savefig(path, transparent=True, bbox_inches="tight", pad_inches=0.02)
    plt.close(fig)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def add_formula(doc: Document, image_path: Path, number: str, *, width: float = 4.6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t")
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.add_run("\t")
    number_run = paragraph.add_run(number)
    set_run_font(number_run)
    doc.add_paragraph()


def add_title_page(doc: Document) -> None:
    lines_top = [
        "Министерство образования и науки Российской Федерации",
        "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        "«Ижевский государственный технический университет имени М. Т. Калашникова»",
    ]
    for line in lines_top:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    for _ in range(6):
        add_paragraph(doc, "", indent=False)

    add_paragraph(doc, "Лабораторная работа N 2", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
    add_paragraph(
        doc,
        "Решение задач линейного программирования",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        bold=True,
    )
    add_paragraph(doc, "По дисциплине «Методы оптимизации»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph(doc, "Вариант 16", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    for _ in range(5):
        add_paragraph(doc, "", indent=False)

    add_paragraph(doc, "Выполнил: студент гр. М25-787-1 Р. В. Скороходов", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    add_paragraph(
        doc,
        "Принял: доктор физико-математических наук, профессор В. А. Тененев",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        indent=False,
    )
    for _ in range(4):
        add_paragraph(doc, "", indent=False)
    add_paragraph(doc, "Ижевск 2026", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    page_break(doc)


def add_contents_page(doc: Document) -> None:
    add_paragraph(doc, "Содержание", bold=True, indent=False)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Содержание будет обновлено автоматически."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    page_break(doc)


def update_word_fields(path: Path) -> None:
    ps_path = str(path).replace("'", "''")
    command = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{ps_path}')
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
$doc.Fields.Update() | Out-Null
$doc.Save()
$doc.Close($false)
$word.Quit()
"""
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", command],
            check=True,
            capture_output=True,
            text=True,
        )
    except Exception as exc:
        print(f"Warning: could not update Word fields automatically: {exc}")


def render_formulas() -> dict[str, Path]:
    formulas = {
        "objective": (["Z=-3x_1+6x_2"], False, 3.2),
        "system": (
            [
                r"5x_1-2x_2 \leq 4,",
                r"x_1-2x_2 \geq -4,",
                r"x_1+x_2 \geq 4,",
                r"x_1,x_2 \geq 0.",
            ],
            True,
            4.5,
        ),
        "level_line": ([r"-3x_1+6x_2=C"], False, 3.0),
        "standard": (
            [
                r"5x_1-2x_2+s_1=4,",
                r"-x_1+2x_2+s_2=4,",
                r"x_1+x_2-e_3+a_3=4.",
            ],
            True,
            4.7,
        ),
        "phase1": ([r"F=-a_3 \to \max"], False, 2.5),
        "min_transform": ([r"\min Z \Longleftrightarrow \max(-Z)=3x_1-6x_2"], False, 4.8),
        "answer": (
            [
                r"x^*_{\min}=\left(\frac{12}{7};\frac{16}{7}\right),\quad Z_{\min}=\frac{60}{7},",
                r"Z_{\max}=12,\quad x \in \left[\left(\frac{4}{3};\frac{8}{3}\right),\left(2;3\right)\right].",
            ],
            False,
            5.6,
        ),
    }
    paths: dict[str, Path] = {}
    for name, (lines, brace, width) in formulas.items():
        path = FORMULA_DIR / f"{name}.png"
        render_math(lines, path, brace=brace, width=width)
        paths[name] = path
    return paths


def vertices_table() -> pd.DataFrame:
    graph = solve_graphically(VARIANT_16)
    rows = []
    for idx, vertex in enumerate(graph.vertices, start=1):
        rows.append(
            {
                "Точка": f"A{idx}",
                "x1": format_fraction(vertex.x1),
                "x2": format_fraction(vertex.x2),
                "Z": format_fraction(vertex.value),
                "Активные ограничения": "; ".join(vertex.active_constraints),
            }
        )
    return pd.DataFrame(rows)


def steps_table(result, phase: str) -> pd.DataFrame:
    rows = []
    steps = result.phase1 if phase == "phase I" else result.phase2
    for step in steps:
        rows.append(
            {
                "k": step.iteration,
                "Базис": ", ".join(step.basis),
                "Свободные члены": "; ".join(as_fraction(value) for value in step.rhs),
                "Целевая строка": as_fraction(step.objective_value),
                "Входит": step.entering or "-",
                "Выходит": step.leaving or "-",
            }
        )
    return pd.DataFrame(rows)


def final_tableau(result) -> pd.DataFrame:
    step = result.phase2[-1]
    names = list(step.reduced_costs.keys()) + ["b"]
    rows = []
    for basis_name, values in zip(step.basis, step.tableau):
        row = {"Базис": basis_name}
        for name, value in zip(names, values):
            row[name] = as_fraction(value)
        rows.append(row)
    return pd.DataFrame(rows)


def add_method_structure(
    doc: Document,
    *,
    title: str,
    idea: str,
    data: str,
    iteration: str,
    stop: str,
    result: str,
    graph: str,
) -> None:
    add_heading(doc, title, level=1)
    items = [
        ("Идея метода.", idea),
        ("Какие данные нужны для запуска.", data),
        ("Как выполняется одна итерация.", iteration),
        ("Какой критерий остановки.", stop),
        ("Что получилось в данной задаче.", result),
        ("Что видно на графике.", graph),
    ]
    for label, text in items:
        add_paragraph(doc, f"{label} {text}")


def main() -> None:
    formulas = render_formulas()
    graph = solve_graphically(VARIANT_16)
    simplex_max = solve_with_simplex(VARIANT_16, "max")
    simplex_min = solve_with_simplex(VARIANT_16, "min")

    doc = Document()
    configure_document(doc)

    add_title_page(doc)
    add_contents_page(doc)

    add_heading(doc, "1 Цель работы")
    add_paragraph(
        doc,
        "Цель работы - решить задачу линейного программирования графическим методом и методом симплекс-таблиц, "
        "а затем сравнить найденные максимум и минимум целевой функции.",
    )

    add_heading(doc, "1.1 Идея работы простыми словами", level=2)
    add_paragraph(
        doc,
        "В задаче нужно выбрать такие значения x1 и x2, которые одновременно удовлетворяют всем ограничениям. "
        "На плоскости эти ограничения образуют допустимую область. После этого проверяется, где линейная функция Z "
        "становится наименьшей и наибольшей. Графический метод показывает это на рисунке, а симплекс-метод делает "
        "то же самое через последовательные таблицы.",
    )

    add_heading(doc, "1.2 Словарь обозначений", level=2)
    glossary = [
        "x = (x1, x2) - точка на плоскости;",
        "Z = f(x1, x2) - значение целевой функции в этой точке;",
        "xk - точка или базисное решение, полученное на k-й итерации;",
        "x* - точка минимума или максимума;",
        "Zmin, Zmax - минимальное и максимальное значения функции;",
        "s - добавочная переменная для ограничения типа <=;",
        "e - избыточная переменная для ограничения типа >=;",
        "a - искусственная переменная, которая нужна для начального базиса;",
        "epsilon - малое число, по которому проверяется остановка вычислений.",
    ]
    for item in glossary:
        add_paragraph(doc, item)

    add_heading(doc, "2 Постановка задачи")
    add_paragraph(
        doc,
        "Зачем нужна следующая формула: она задает показатель, который требуется сначала минимизировать, "
        "а затем максимизировать при одних и тех же ограничениях.",
    )
    add_formula(doc, formulas["objective"], "(2.1)", width=3.0)
    add_paragraph(
        doc,
        "Зачем нужна следующая система: она задает область допустимых решений. Все точки вне этой области "
        "в задаче рассматривать нельзя.",
    )
    add_formula(doc, formulas["system"], "(2.2)", width=3.9)

    add_method_structure(
        doc,
        title="3 Графический метод",
        idea="Каждое ограничение строится как полуплоскость. Их пересечение дает допустимую область.",
        data="Нужны коэффициенты ограничений, правая часть каждого ограничения и коэффициенты целевой функции.",
        iteration="Для двухмерной задачи достаточно найти точки пересечения граничных прямых и оставить только допустимые точки.",
        stop="Остановка происходит после проверки всех пересечений. Дополнительных итераций здесь нет.",
        result=(
            "Найдены три вершины допустимого треугольника: (4/3; 8/3), (12/7; 16/7) и (2; 3). "
            "Минимум достигается в точке (12/7; 16/7), максимум - на отрезке между (4/3; 8/3) и (2; 3)."
        ),
        graph=(
            "На графике видно, что линия уровня целевой функции параллельна одному ребру допустимой области. "
            "Поэтому максимум получается не в одной точке, а на целом отрезке."
        ),
    )
    add_paragraph(
        doc,
        "Зачем нужна следующая формула: она показывает семейство линий уровня. При разных C линия двигается "
        "параллельно самой себе, и по последнему касанию с допустимой областью находится экстремум.",
    )
    add_formula(doc, formulas["level_line"], "(3.1)", width=2.8)

    vertices = vertices_table()
    add_table(doc, vertices, "Таблица 1 - Значения целевой функции в вершинах допустимой области", size=8)
    add_paragraph(
        doc,
        "Что я должен понять из таблицы: значение Z минимально в точке A2, где Z = 60/7. "
        "В точках A1 и A3 значение одинаковое и равно 12, поэтому максимум достигается на всем ребре между этими точками. "
        "Это важная проверка: если бы смотреть только на одну вершину, можно было бы пропустить неединственность максимума.",
    )

    figure_path = ROOT / "figures" / "feasible_region.png"
    if figure_path.exists():
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(figure_path), width=Inches(5.5))
        caption = doc.add_paragraph("Рисунок 1 - Допустимая область и точки экстремума")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_method_structure(
        doc,
        title="4 Симплекс-метод для максимума",
        idea="Метод переходит от одного базисного допустимого решения к другому так, чтобы значение целевой функции не уменьшалось.",
        data="Нужны ограничения в канонической форме, начальный базис и коэффициенты функции Z.",
        iteration="На каждой итерации выбирается входящая переменная с положительной оценкой и выходящая переменная по минимальному отношению свободного члена к положительному коэффициенту.",
        stop="Метод останавливается, когда в строке оценок больше нет положительных коэффициентов.",
        result="Получено Zmax = 12. Так как в оптимальной таблице есть нулевая оценка у небазисной переменной, максимум не единственный.",
        graph="График подтверждает этот вывод: оптимальные точки лежат на ребре от (4/3; 8/3) до (2; 3).",
    )
    add_paragraph(
        doc,
        "Зачем нужна следующая формула: ограничения нужно привести к равенствам, потому что симплекс-таблица работает "
        "с базисными переменными. Второе ограничение предварительно умножено на -1, чтобы правая часть стала положительной.",
    )
    add_formula(doc, formulas["standard"], "(4.1)", width=4.2)
    add_paragraph(
        doc,
        "Зачем нужна следующая формула: фаза I убирает искусственную переменную из базиса и проверяет, существует ли "
        "допустимое начальное решение.",
    )
    add_formula(doc, formulas["phase1"], "(4.2)", width=2.3)

    add_table(doc, steps_table(simplex_max, "phase I"), "Таблица 2 - Фаза I симплекс-метода", size=8)
    add_paragraph(
        doc,
        "Что я должен понять из таблицы: сначала искусственная переменная a3 находится в базисе, потому что без нее "
        "начальную таблицу построить неудобно. После двух переходов значение вспомогательной функции становится равным нулю. "
        "Значит, допустимое решение найдено и можно переходить к исходной целевой функции.",
    )
    add_table(doc, steps_table(simplex_max, "phase II"), "Таблица 3 - Фаза II для максимума", size=8)
    add_paragraph(
        doc,
        "Что я должен понять из таблицы: на первой строке фазы II значение Z еще равно 60/7, то есть это не максимум. "
        "После ввода переменной e3 получается значение 12. Положительных оценок больше нет, поэтому симплекс-метод останавливается.",
    )
    add_table(doc, final_tableau(simplex_max), "Таблица 4 - Итоговая симплекс-таблица для максимума", size=8)
    add_paragraph(
        doc,
        "Что я должен понять из таблицы: итоговый базис дает точку (2; 3), но нулевая оценка у небазисной переменной "
        "показывает альтернативный оптимум. Поэтому правильный ответ для максимума - не одна точка, а отрезок.",
    )

    add_method_structure(
        doc,
        title="5 Симплекс-метод для минимума",
        idea="Минимизацию удобно свести к максимизации противоположной функции -Z.",
        data="Используются те же ограничения и тот же допустимый базис, который был получен на фазе I.",
        iteration="Итерация выполняется так же, как для максимума: проверяются оценки и при необходимости меняется базис.",
        stop="Остановка происходит, когда для функции -Z нет положительных оценок.",
        result="Фаза II сразу останавливается в точке (12/7; 16/7), где исходная функция имеет значение 60/7.",
        graph="На графике эта точка является нижней вершиной допустимого треугольника относительно направления убывания Z.",
    )
    add_paragraph(
        doc,
        "Зачем нужна следующая формула: симплекс-алгоритм в работе реализован для максимизации, поэтому минимум исходной "
        "функции находится через максимум функции с противоположным знаком.",
    )
    add_formula(doc, formulas["min_transform"], "(5.1)", width=4.5)
    add_table(doc, steps_table(simplex_min, "phase II"), "Таблица 5 - Фаза II для минимума через максимум -Z", size=8)
    add_paragraph(
        doc,
        "Что я должен понять из таблицы: после фазы I уже получена вершина (12/7; 16/7). "
        "Для функции -Z положительных оценок нет, поэтому двигаться в соседнюю вершину не нужно. "
        "Значение исходной функции в этой точке равно 60/7.",
    )
    add_table(doc, final_tableau(simplex_min), "Таблица 6 - Итоговая симплекс-таблица для минимума", size=8)
    add_paragraph(
        doc,
        "Что я должен понять из таблицы: базис содержит x1 и x2, поэтому координаты решения читаются прямо из столбца "
        "свободных членов. Остальные переменные показывают запас или избыточность по ограничениям и не меняют найденную точку.",
    )

    add_heading(doc, "6 Контрольная проверка")
    check = pd.DataFrame(
        [
            {
                "Метод": "Графический",
                "Минимум": "(12/7; 16/7), Z = 60/7",
                "Максимум": "отрезок [(4/3; 8/3), (2; 3)], Z = 12",
            },
            {
                "Метод": "Симплекс-таблицы",
                "Минимум": "(12/7; 16/7), Z = 60/7",
                "Максимум": "(2; 3), Z = 12 и признак альтернативного оптимума",
            },
        ]
    )
    add_table(doc, check, "Таблица 7 - Сравнение результатов", size=8)
    add_paragraph(
        doc,
        "Что я должен понять из таблицы: оба метода дают одинаковое минимальное значение и одинаковое максимальное "
        "значение. Отличие только в форме записи максимума: графический метод сразу показывает весь отрезок, а симплекс "
        "дает одну вершину и дополнительный признак альтернативного оптимума.",
    )

    add_heading(doc, "7 Вывод")
    add_paragraph(
        doc,
        "В работе я последовательно проверил задачу двумя способами. Сначала по графику были найдены все вершины "
        "допустимой области и значения целевой функции в них. Затем та же задача была решена через двухфазный "
        "симплекс-метод. Результаты совпали: минимум равен 60/7 в точке (12/7; 16/7), максимум равен 12 на отрезке "
        "от (4/3; 8/3) до (2; 3).",
    )
    add_paragraph(
        doc,
        "Зачем нужна итоговая формула: она компактно фиксирует окончательный ответ, который дальше можно сравнивать "
        "с графиком и симплекс-таблицами.",
    )
    add_formula(doc, formulas["answer"], "(7.1)", width=5.3)

    add_heading(doc, "8 Список использованных источников")
    add_paragraph(
        doc,
        "Лабораторная работа N 2. Решение задач линейного программирования: методические указания.",
        indent=False,
    )
    add_paragraph(
        doc,
        "Репозиторий с исходным кодом и ноутбуками: https://github.com/megusto0/mp-2.",
        indent=False,
    )

    doc.save(REPORT_PATH)
    update_word_fields(REPORT_PATH)
    print(f"Saved {REPORT_PATH}")


if __name__ == "__main__":
    main()
